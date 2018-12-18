import itertools

import attr
import docx
import docx.enum.section
import docx.enum.text

import epyqlib.cangenmanual
import epyqlib.pm.parametermodel
import epyqlib.treenode
import epyqlib.utils.general

builders = epyqlib.utils.general.TypeMap()


@attr.s
class Row:
    name = attr.ib()
    indent = attr.ib(default=0)
    fill = attr.ib(default='')
    factor = attr.ib(default='')
    units = attr.ib(default='')
    default = attr.ib(default='')
    minimum = attr.ib(default='')
    maximum = attr.ib(default='')
    enumeration = attr.ib(default='')
    comment = attr.ib(default='')

    def to_tuple(self, max_indent=5):
        return (
            *((self.fill,) * self.indent),
            self.name,
            *((self.fill,) * (max_indent - self.indent - 1)),
            self.factor,
            self.units,
            self.default,
            self.minimum,
            self.maximum,
            self.enumeration,
            self.comment,
        )


@builders(epyqlib.pm.parametermodel.Root)
@attr.s
class Root:
    wrapped = attr.ib()
    can_root = attr.ib()
    template = attr.ib()
    access_level = attr.ib()

    def gen(self):
        headings = Row(
            name='Name',
            factor='Factor',
            units='Units',
            default='Default',
            minimum='Min',
            maximum='Max',
            enumeration='Enumeration',
            comment='Comment',
        )
        max_indent = 5
        fill_width = 0.25
        name_width = 2.75
        widths = Row(
            name=name_width - (fill_width * (max_indent - 1)),
            indent=max_indent - 1,
            fill=fill_width,
            factor=0.625,
            units=0.625,
            default=0.875,
            minimum=0.625,
            maximum=0.625,
            enumeration=1.5,
            comment=None,
        )

        for heading, width in zip(headings.to_tuple(), widths.to_tuple()):
            print(heading, width)

        table = epyqlib.cangenmanual.Table(
            title=None,
            headings=headings.to_tuple(),
            widths=widths.to_tuple(),
            total_width=10,
        )

        import time
        start = time.monotonic()

        for child in self.wrapped.children:
            if child.name.endswith('Other'):
                continue

            try:
                builder = builders.wrap(
                    wrapped=child,
                    parameter_root=self.wrapped,
                    can_root=self.can_root,
                    access_level=self.access_level,
                )
            except KeyError:
                continue

            table.rows.extend(builder.gen(indent=0))

        now = time.monotonic()
        delta = now - start
        start = now
        print('rows built', int(delta))

        raw_rows = table.rows

        table.rows = tuple(
            row.to_tuple(max_indent=max_indent)
            for row in table.rows
        )

        now = time.monotonic()
        delta = now - start
        start = now
        print('rows converted', int(delta))

        if self.template is not None:
            doc = docx.Document(self.template)
        else:
            doc = docx.Document()

        doc_table = doc.add_table(rows=0, cols=len(table.headings))
        doc_table.autofit = False

        table.fill_docx(doc_table)

        epyqlib.cangenmanual.set_repeat_table_header(doc_table.rows[0])
        epyqlib.cangenmanual.prevent_row_breaks(doc_table)

        for row, doc_row in zip(itertools.chain((headings,), raw_rows), doc_table.rows):
            base_cell = doc_row.cells[row.indent]
            base_cell.merge(doc_row.cells[max_indent - 1])

        for row in doc_table.rows:
            for cell in row.cells:
                cell.text = cell.text.strip()
                cell.paragraphs[0].paragraph_format.line_spacing_rule = (
                    docx.enum.text.WD_LINE_SPACING.SINGLE
                )

        now = time.monotonic()
        delta = now - start
        start = now
        print('table filled', int(delta))

        for section in doc.sections:
            section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE

        return doc


@builders(epyqlib.pm.parametermodel.Group)
@attr.s
class Group:
    wrapped = attr.ib()
    parameter_root = attr.ib()
    can_root = attr.ib()
    access_level = attr.ib()

    def gen(self, indent):
        rows = [
            Row(
                name=self.wrapped.name,
                indent=indent,
            ),
        ]

        for child in self.wrapped.children:
            builder = builders.wrap(
                wrapped=child,
                parameter_root=self.parameter_root,
                can_root=self.can_root,
                access_level=self.access_level,
            )
            rows.extend(builder.gen(indent=indent + 1))

        return rows


@builders(epyqlib.pm.parametermodel.Parameter)
@attr.s
class Parameter:
    wrapped = attr.ib()
    parameter_root = attr.ib()
    can_root = attr.ib()
    access_level = attr.ib()

    def gen(self, indent):
        try:
            access_level = self.parameter_root.nodes_by_attribute(
                attribute_value=self.wrapped.access_level_uuid,
                attribute_name='uuid',
            ).pop()
        except epyqlib.treenode.NotFoundError:
            pass
        else:
            if access_level.value > self.access_level.value:
                print('skipping', self.wrapped.name)
                return []

        signal = self.can_root.nodes_by_attribute(
            attribute_value=self.wrapped.uuid,
            attribute_name='parameter_uuid',
        ).pop()

        factor = signal.factor
        if factor is None or factor == 1:
            factor = ''

        units = self.wrapped.units
        if units is None:
            units = ''

        try:
            enumeration = self.parameter_root.nodes_by_attribute(
                attribute_value=self.wrapped.enumeration_uuid,
                attribute_name='uuid',
            ).pop().name
        except epyqlib.treenode.NotFoundError:
            enumeration = ''

        default = self.wrapped.default
        if default is None:
            default = ''

        minimum = self.wrapped.minimum
        if minimum is None:
            minimum = ''

        maximum = self.wrapped.maximum
        if maximum is None:
            maximum = ''

        comment = self.wrapped.comment
        if comment is None:
            comment = ''

        return [
            Row(
                name=self.wrapped.name,
                indent=indent,
                factor=factor,
                units=units,
                default=default,
                minimum=minimum,
                maximum=maximum,
                enumeration=enumeration,
                comment=comment,
            ),
        ]
